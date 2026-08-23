#!/usr/bin/env python
"""Generate every CV output from cv-data.yaml.

    python build.py                      # everything
    python build.py --profile resume     # only the 2-page resume
    python build.py --lang en --format docx

Nothing generated is ever hand-edited: edit cv-data.yaml and rebuild. That
includes the two published web pages — ../index.html and ./index.html — which
are as generated as anything in out/.
"""
from __future__ import annotations

import argparse
import datetime as dt
import json
import math
import re
import sys
from pathlib import Path

import yaml

ROOT = Path(__file__).parent
DATA = ROOT / "cv-data.yaml"
OUT = ROOT / "out"
SITE = ROOT.parent               # repo root — GitHub Pages serves from here
LANDING = SITE / "index.html"    #   →  kristonova.github.io/
WEB_CV = ROOT / "index.html"     #   →  kristonova.github.io/cv/

LANGS = ("en", "id")
PROFILES = {"resume": 1, "academic": 2}  # max tier included

MONTHS = {
    "en": "Jan Feb Mar Apr May Jun Jul Aug Sep Oct Nov Dec".split(),
    "id": "Jan Feb Mar Apr Mei Jun Jul Agu Sep Okt Nov Des".split(),
}
MONTHS_LONG = {
    "en": ["January", "February", "March", "April", "May", "June", "July",
           "August", "September", "October", "November", "December"],
    "id": ["Januari", "Februari", "Maret", "April", "Mei", "Juni", "Juli",
           "Agustus", "September", "Oktober", "November", "Desember"],
}

L = {
    "present": {"en": "Present", "id": "Saat ini"},
    "graduated": {"en": "Graduated", "id": "Lulus"},
    "expired": {"en": "expired", "id": "kedaluwarsa"},
    "issued": {"en": "issued", "id": "terbit"},
    "reg": {"en": "Reg.", "id": "No. Reg."},
    "resume_note": {"en": "Résumé", "id": "Résumé"},
    # section titles
    "summary": {"en": "Profile", "id": "Profil Profesional"},
    "education": {"en": "Education", "id": "Pendidikan"},
    "research": {"en": "Research Experience", "id": "Pengalaman Riset"},
    "publications": {"en": "Publications", "id": "Publikasi Ilmiah"},
    "experience": {"en": "Professional Experience", "id": "Pengalaman Profesional"},
    "teaching": {"en": "Teaching Experience", "id": "Pengalaman Mengajar"},
    "talks": {"en": "Invited Talks, Lectures & Seminars", "id": "Pembicara, Kuliah Umum & Seminar"},
    "certifications": {"en": "Professional Certifications", "id": "Sertifikasi Profesi"},
    "training": {"en": "Professional Training", "id": "Pelatihan Profesional"},
    "community": {"en": "Community Service & Leadership", "id": "Organisasi & Pengabdian Masyarakat"},
    "skills": {"en": "Skills", "id": "Keahlian Teknis"},
    "languages": {"en": "Languages", "id": "Kemampuan Bahasa"},
    "awards": {"en": "Honours & Awards", "id": "Penghargaan & Beasiswa"},
    "engagements": {"en": "Selected client engagements", "id": "Proyek & Klien Terpilih"},
}

TALK_KIND = {
    "paper":           {"en": "Paper Presentation", "id": "Presentasi Makalah Ilmiah"},
    "general_lecture": {"en": "General Lecture", "id": "Kuliah Umum"},
    "guest_lecture":   {"en": "Guest Lecture", "id": "Kuliah Tamu"},
    "seminar":         {"en": "Seminar", "id": "Seminar"},
    "webinar":         {"en": "Webinar", "id": "Webinar"},
    "workshop":        {"en": "Workshop", "id": "Workshop / Lokakarya"},
    "fgd":             {"en": "Focus Group Discussion", "id": "Focus Group Discussion (FGD)"},
}

# Section order per profile.
ORDER = {
    "resume": ["summary", "experience", "education", "skills", "languages",
               "publications", "certifications", "community", "awards"],
    "academic": ["summary", "education", "research", "publications", "experience",
                 "teaching", "talks", "certifications", "training", "community",
                 "skills", "languages", "awards"],
}


# --------------------------------------------------------------------------- #
# helpers
# --------------------------------------------------------------------------- #
def pick(value, lang: str, default: str = "") -> str:
    """Resolve a {en:..., id:...} mapping, a plain scalar, or None."""
    if value is None:
        return default
    if isinstance(value, dict):
        return value.get(lang) or value.get("en") or default
    return str(value)


def pick_list(value, lang: str) -> list[str]:
    if not value:
        return []
    if isinstance(value, dict):
        return list(value.get(lang) or value.get("en") or [])
    return list(value)


def fmt_point(value, lang: str) -> str:
    """Format YYYY, 'YYYY-MM', or a date into 'Mon YYYY'."""
    if value is None:
        return L["present"][lang]
    if isinstance(value, (dt.date, dt.datetime)):
        return f"{MONTHS[lang][value.month - 1]} {value.year}"
    s = str(value)
    m = re.fullmatch(r"(\d{4})-(\d{2})", s)
    if m:
        return f"{MONTHS[lang][int(m.group(2)) - 1]} {m.group(1)}"
    return s


def fmt_range(start, end, lang: str) -> str:
    a, b = fmt_point(start, lang), fmt_point(end, lang)
    return a if a == b else f"{a} – {b}"


def fmt_long(value, lang: str) -> str:
    """Full date, e.g. '29 July 2026' — used for the graduation line."""
    if isinstance(value, (dt.date, dt.datetime)):
        return f"{value.day} {MONTHS_LONG[lang][value.month - 1]} {value.year}"
    return fmt_point(value, lang)


def has_day(value) -> bool:
    return isinstance(value, (dt.date, dt.datetime))


def run(text: str, b: bool = False, i: bool = False) -> dict:
    return {"t": text, "b": b, "i": i}


def bilingual(fn) -> dict:
    """Call fn(lang) for each language and collect into {en:..., id:...}."""
    return {lang: fn(lang) for lang in LANGS}


# --------------------------------------------------------------------------- #
# document model  (one structure, three renderers)
# --------------------------------------------------------------------------- #
def block(tier, primary, meta=None, sub=None, bullets=None, links=None,
          children=None, group_label=None) -> dict:
    return {"style": "block", "tier": tier, "primary": primary, "meta": meta,
            "sub": sub, "bullets": bullets or {}, "links": links or [],
            "children": children or [], "group_label": group_label}


def line(tier, runs, meta=None, links=None) -> dict:
    return {"style": "line", "tier": tier, "runs": runs, "meta": meta,
            "links": links or []}


def build_model(d: dict) -> dict:
    """Turn the raw YAML into renderer-agnostic sections."""
    sections: dict[str, list] = {}

    # --- summary -----------------------------------------------------------
    sections["summary"] = [line(1, bilingual(
        lambda lang: [run(" ".join(pick(d["summary"], lang).split()))]))]

    # --- education ---------------------------------------------------------
    entries = []
    for e in d["education"]:
        details = bilingual(lambda lang, e=e: (
            ([f"{L['graduated'][lang]} {fmt_long(e['end'], lang)}."] if has_day(e["end"]) else [])
            + pick_list(e.get("details"), lang)))
        entries.append(block(
            e["tier"],
            bilingual(lambda lang, e=e: f"{pick(e['degree'], lang)} — {e['institution']}"),
            meta=bilingual(lambda lang, e=e: fmt_range(e["start"], e["end"], lang)),
            sub=bilingual(lambda lang, e=e: pick(e.get("location"), lang)),
            bullets=details))
    sections["education"] = entries

    # --- research / teaching / community (same shape) ----------------------
    for key, name_field in (("research", "org"), ("teaching", "org"), ("community", "org")):
        entries = []
        for e in d.get(key, []):
            if e.get("archived"):
                continue
            title = e.get("title")
            head = bilingual(lambda lang, e=e, title=title: (
                f"{pick(title, lang)} — {e['org']}" if title
                else f"{e['org']} — {pick(e['role'], lang)}"))
            sub = bilingual(lambda lang, e=e, title=title: (
                pick(e.get("role"), lang) if title else pick(e.get("location"), lang)))
            entries.append(block(
                e["tier"], head,
                meta=bilingual(lambda lang, e=e: fmt_range(e["start"], e["end"], lang)),
                sub=sub,
                bullets=bilingual(lambda lang, e=e: pick_list(e.get("bullets"), lang))))
        sections[key] = entries

    # --- experience (with nested client engagements) -----------------------
    entries = []
    for e in d["experience"]:
        if e.get("archived"):
            continue
        children = []
        for g in e.get("engagements", []):
            if g.get("archived"):
                continue
            children.append(block(
                g["tier"],
                bilingual(lambda lang, g=g: f"{g['client']} — {pick(g['role'], lang)}"),
                meta=bilingual(lambda lang, g=g: fmt_range(g["start"], g["end"], lang)),
                bullets=bilingual(lambda lang, g=g: pick_list(g.get("bullets"), lang)),
                links=g.get("links", [])))
        entries.append(block(
            e["tier"],
            bilingual(lambda lang, e=e: f"{e['org']} — {pick(e['role'], lang)}"),
            meta=bilingual(lambda lang, e=e: fmt_range(e["start"], e["end"], lang)),
            sub=bilingual(lambda lang, e=e: pick(e.get("location"), lang)),
            bullets=bilingual(lambda lang, e=e: pick_list(e.get("bullets"), lang)),
            links=e.get("links", []),
            children=children,
            group_label=L["engagements"]))
    sections["experience"] = entries

    # --- publications ------------------------------------------------------
    entries = []
    for p in d["publications"]:
        citation = " ".join(p["citation"].split())
        tail = f" DOI: {p['doi']}." if p.get("doi") else ""
        entries.append(line(p["tier"], bilingual(
            lambda lang, c=citation, t=tail: [run(c + t)]), links=[p["url"]] if p.get("url") else []))
    sections["publications"] = entries

    # --- talks -------------------------------------------------------------
    entries = []
    for t in d["talks"]:
        entries.append(line(
            t["tier"],
            bilingual(lambda lang, t=t: [
                run(pick(TALK_KIND[t["kind"]], lang), b=True),
                run(f", “{t['title']},” "),
                run(t["venue"], i=True),
                run("."),
            ]),
            meta=bilingual(lambda lang, t=t: fmt_point(t["date"], lang)),
            links=[t["url"]] if t.get("url") else []))
    sections["talks"] = entries

    # --- certifications ----------------------------------------------------
    entries = []
    for c in d["certifications"]:
        def runs(lang, c=c):
            issued = fmt_point(c["issued"], lang)
            status = f", {L['expired'][lang]}" if c.get("expired") else ""
            r = [run(pick(c["name"], lang), b=True),
                 run(f" — {c['issuer']}, {L['reg'][lang]} {c['reg']}"),
                 run(f" ({L['issued'][lang]} {issued}, {c['valid_years']}y{status})")]
            note = pick(c.get("note"), lang)
            if note:
                r.append(run(f" {note}", i=True))
            return r
        entries.append(line(c["tier"], bilingual(runs)))
    sections["certifications"] = entries

    # --- training ----------------------------------------------------------
    entries = []
    for t in d.get("training", []):
        entries.append(line(
            t["tier"],
            bilingual(lambda lang, t=t: [run(t["title"], b=True), run(f" — {t['org']}.")]),
            meta=bilingual(lambda lang, t=t: fmt_point(t["date"], lang))))
    sections["training"] = entries

    # --- skills ------------------------------------------------------------
    entries = []
    for s in d["skills"]:
        entries.append(line(s["tier"], bilingual(
            lambda lang, s=s: [run(f"{pick(s['group'], lang)}: ", b=True),
                               run(", ".join(s["items"]))])))
    sections["skills"] = entries

    # --- languages ---------------------------------------------------------
    entries = []
    for lg in d["languages"]:
        def runs(lang, lg=lg):
            r = [run(f"{pick(lg['name'], lang)}: ", b=True), run(pick(lg["level"], lang))]
            if lg.get("detail"):
                r.append(run(f" — {lg['detail']}"))
            return r
        entries.append(line(lg["tier"], bilingual(runs)))
    sections["languages"] = entries

    # --- awards ------------------------------------------------------------
    entries = []
    for a in d["awards"]:
        def runs(lang, a=a):
            r = [run(pick(a["name"], lang), b=True), run(f" — {pick(a['org'], lang)}")]
            note = pick(a.get("note"), lang)
            if note:
                r.append(run(f". {note}", i=True))
            return r
        span = f"{a['start']}" if a["start"] == a["end"] else f"{a['start']}–{a['end']}"
        entries.append(line(a["tier"], bilingual(runs),
                            meta=bilingual(lambda lang, span=span: span)))
    sections["awards"] = entries

    return sections


def visible(entries: list, max_tier: int) -> list:
    out = []
    for e in entries:
        if e["tier"] > max_tier:
            continue
        e = dict(e)
        if e.get("children"):
            e["children"] = [c for c in e["children"] if c["tier"] <= max_tier]
        out.append(e)
    return out


# --------------------------------------------------------------------------- #
# Markdown
# --------------------------------------------------------------------------- #
def md_runs(runs: list) -> str:
    """Emphasis markers must hug the text — Markdown ignores `** text **`."""
    parts = []
    for r in runs:
        core = r["t"].strip()
        if not core:
            parts.append(r["t"])
            continue
        lead = r["t"][: len(r["t"]) - len(r["t"].lstrip())]
        trail = r["t"][len(r["t"].rstrip()):]
        if r["b"]:
            core = f"**{core}**"
        if r["i"]:
            core = f"*{core}*"
        parts.append(f"{lead}{core}{trail}")
    return "".join(parts)


def render_markdown(meta: dict, sections: dict, lang: str, profile: str) -> str:
    max_tier = PROFILES[profile]
    o: list[str] = [f"# {meta['name']}", ""]
    o.append(f"**{pick(meta['headline'], lang)}**")
    o.append("")
    contact = [pick(meta["location"], lang)]
    if meta.get("phone"):
        contact.append(meta["phone"])
    contact.append(meta["email"])
    o.append(" · ".join(contact))
    o.append("")
    o.append(" · ".join(f"[{k}]({v})" for k, v in meta["links"].items()))
    o.append("")

    for key in ORDER[profile]:
        entries = visible(sections.get(key, []), max_tier)
        if not entries:
            continue
        o.append(f"## {L[key][lang]}")
        o.append("")
        for e in entries:
            if e["style"] == "line":
                meta_txt = pick(e.get("meta"), lang)
                text = md_runs(e["runs"][lang])
                if key == "summary":
                    o.extend([text, ""])
                    continue
                links = "".join(f" <{u}>" for u in e["links"])
                suffix = f" ({meta_txt})" if meta_txt else ""
                o.append(f"- {text}{suffix}{links}")
            else:
                o.append(f"**{pick(e['primary'], lang)}** — {pick(e.get('meta'), lang)}")
                sub = pick(e.get("sub"), lang)
                if sub:
                    o.append(f"*{sub}*")
                o.append("")
                for b in e["bullets"].get(lang, []):
                    o.append(f"- {b}")
                for u in e["links"]:
                    o.append(f"- <{u}>")
                if e["bullets"].get(lang) or e["links"]:
                    o.append("")
                if e["children"]:
                    o.append(f"*{pick(e['group_label'], lang)}:*")
                    o.append("")
                    for c in e["children"]:
                        o.append(f"- **{pick(c['primary'], lang)}** "
                                 f"({pick(c.get('meta'), lang)})")
                        for b in c["bullets"].get(lang, []):
                            o.append(f"  - {b}")
                        for u in c["links"]:
                            o.append(f"  - <{u}>")
                    o.append("")
        if o[-1] != "":
            o.append("")
    return "\n".join(o).rstrip() + "\n"


# --------------------------------------------------------------------------- #
# DOCX  (single column, built-in heading styles, no tables — ATS-safe)
# --------------------------------------------------------------------------- #
def render_docx(meta: dict, sections: dict, lang: str, profile: str, path: Path) -> None:
    from docx import Document
    from docx.enum.text import WD_TAB_ALIGNMENT
    from docx.shared import Inches, Pt, RGBColor

    max_tier = PROFILES[profile]
    body_pt = 9.5 if profile == "resume" else 10.5
    doc = Document()

    for s in doc.sections:
        s.left_margin = s.right_margin = Inches(0.65)
        s.top_margin = s.bottom_margin = Inches(0.55)
    content_width = Inches(8.5 - 0.65 * 2)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(body_pt)
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = 1.0

    def para(style=None, space_before=0, space_after=2):
        p = doc.add_paragraph(style=style)
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        return p

    def tabbed(space_before=0, space_after=1):
        p = para(space_before=space_before, space_after=space_after)
        p.paragraph_format.tab_stops.add_tab_stop(content_width, WD_TAB_ALIGNMENT.RIGHT)
        return p

    def add_runs(p, runs):
        for r in runs:
            rr = p.add_run(r["t"])
            rr.bold, rr.italic = r["b"], r["i"]

    # header
    p = para(space_after=1)
    r = p.add_run(meta["name"])
    r.bold, r.font.size = True, Pt(19)
    p = para(space_after=3)
    r = p.add_run(pick(meta["headline"], lang))
    r.font.size, r.font.color.rgb = Pt(body_pt + 1.5), RGBColor(0x33, 0x33, 0x33)
    p = para(space_after=1)
    contact = [pick(meta["location"], lang)]
    if meta.get("phone"):
        contact.append(meta["phone"])
    contact.append(meta["email"])
    p.add_run(" | ".join(contact)).font.size = Pt(body_pt)
    p = para(space_after=6)
    p.add_run(" | ".join(f"{k}: {v}" for k, v in meta["links"].items())).font.size = Pt(body_pt - 0.5)

    for key in ORDER[profile]:
        entries = visible(sections.get(key, []), max_tier)
        if not entries:
            continue
        h = doc.add_heading(L[key][lang], level=1)
        h.paragraph_format.space_before = Pt(8)
        h.paragraph_format.space_after = Pt(2)
        for hr in h.runs:
            hr.font.size, hr.font.color.rgb = Pt(body_pt + 2), RGBColor(0x1a, 0x3c, 0x6e)

        for e in entries:
            if e["style"] == "line":
                if key == "summary":
                    add_runs(para(space_after=3), e["runs"][lang])
                    continue
                p = para(style="List Bullet", space_after=1)
                add_runs(p, e["runs"][lang])
                m = pick(e.get("meta"), lang)
                if m:
                    p.add_run(f" ({m})")
                for u in e["links"]:
                    p.add_run(f" {u}").font.size = Pt(body_pt - 1)
                continue

            p = tabbed(space_before=4)
            p.add_run(pick(e["primary"], lang)).bold = True
            p.add_run("\t" + pick(e.get("meta"), lang))
            sub = pick(e.get("sub"), lang)
            if sub:
                sp = para(space_after=1)
                sr = sp.add_run(sub)
                sr.italic, sr.font.size = True, Pt(body_pt - 0.5)
            for b in e["bullets"].get(lang, []):
                para(style="List Bullet", space_after=1).add_run(b)
            for u in e["links"]:
                para(style="List Bullet", space_after=1).add_run(u).font.size = Pt(body_pt - 1)

            if e["children"]:
                gp = para(space_before=3, space_after=1)
                gr = gp.add_run(pick(e["group_label"], lang) + ":")
                gr.italic = True
                for c in e["children"]:
                    cp = tabbed(space_before=2)
                    cp.paragraph_format.left_indent = Inches(0.18)
                    cp.add_run(pick(c["primary"], lang)).bold = True
                    cp.add_run("\t" + pick(c.get("meta"), lang))
                    for b in c["bullets"].get(lang, []):
                        bp = para(style="List Bullet", space_after=1)
                        bp.paragraph_format.left_indent = Inches(0.48)
                        bp.add_run(b)
                    for u in c["links"]:
                        bp = para(style="List Bullet", space_after=1)
                        bp.paragraph_format.left_indent = Inches(0.48)
                        bp.add_run(u).font.size = Pt(body_pt - 1)

    doc.save(path)


# --------------------------------------------------------------------------- #
# Landing page  (kristonova.github.io/ — ID first, evidence-dense)
# --------------------------------------------------------------------------- #
def _sort_key(value) -> str:
    """Order mixed 'YYYY-MM' strings and real dates on one axis."""
    return value.isoformat() if isinstance(value, (dt.date, dt.datetime)) else str(value)


def landing_clients(d: dict) -> list[str]:
    """Institution names for the trust strip, most recent first, deduplicated.

    Recognisable institutions are the strongest signal for a non-technical
    Indonesian client, so this is generated from the record rather than curated
    — it cannot drift from the CV. `short:` supplies a display name where the
    formal one is too long to scan.
    """
    rows: list[tuple[int, str, str]] = []
    for e in d["experience"]:
        if e.get("archived"):
            continue
        # Client engagements only. An employer is not a client, and leading the
        # strip with an internship would spend the strongest position on the
        # weakest signal.
        for g in e.get("engagements", []):
            if g.get("archived") or g["tier"] > 2:
                continue
            rows.append((g["tier"], _sort_key(g["start"]), g.get("short") or g["client"]))
    # Tier before date: tier 1 is the significance ranking already in the data,
    # so the names a client recognises fastest land first without inventing a
    # second, hand-curated ordering that could drift. Two stable passes rather
    # than one clever key.
    rows.sort(key=lambda r: r[1], reverse=True)   # newest first
    rows.sort(key=lambda r: r[0])                 # then tier 1 above tier 2

    seen, out = set(), []
    for _, _, name in rows:
        if name not in seen:
            seen.add(name)
            out.append(name)
    return out


def landing_work(d: dict, limit: int = 6) -> list[dict]:
    """Tier-1 engagements as entry blocks — client first, role beneath."""
    items: list[tuple[str, dict]] = []
    for e in d["experience"]:
        if e.get("archived"):
            continue
        if not e.get("engagements") and e["tier"] <= 1:
            items.append((_sort_key(e["start"]), block(
                e["tier"],
                bilingual(lambda lang, e=e: e["org"]),
                meta=bilingual(lambda lang, e=e: fmt_range(e["start"], e["end"], lang)),
                sub=bilingual(lambda lang, e=e: pick(e["role"], lang)),
                bullets=bilingual(lambda lang, e=e: pick_list(e.get("bullets"), lang)),
                links=e.get("links", []))))
        for g in e.get("engagements", []):
            if g.get("archived") or g["tier"] > 1:
                continue
            items.append((_sort_key(g["start"]), block(
                g["tier"],
                bilingual(lambda lang, g=g: g.get("short") or g["client"]),
                meta=bilingual(lambda lang, g=g: fmt_range(g["start"], g["end"], lang)),
                sub=bilingual(lambda lang, g=g: pick(g["role"], lang)),
                bullets=bilingual(lambda lang, g=g: pick_list(g.get("bullets"), lang)),
                links=g.get("links", []))))
    items.sort(key=lambda x: x[0], reverse=True)
    return [b for _, b in items[:limit]]


def landing_stats(d: dict, clients: list[str]) -> dict:
    """Every figure on the page, counted from the record.

    Hardcoded statistics are how a CV page starts lying: the old hand-written
    page claimed "5+ Seminar Nasional" against 44 recorded talks.
    """
    return {
        "talks": str(len(d.get("talks", []))),
        "teaching": str(len([t for t in d.get("teaching", []) if not t.get("archived")])),
        "clients": str(len(clients)),
        "people": f"{d['site']['people_trained']}+",
    }


def landing_jsonld(d: dict) -> str:
    meta, site = d["meta"], d["site"]
    payload = {
        "@context": "https://schema.org",
        "@type": "Person",
        "name": meta["name"],
        "jobTitle": site["role"]["en"],
        "description": " ".join(site["positioning"]["en"].split()),
        "email": f"mailto:{meta['email']}",
        "url": site["url"],
        "image": site["url"] + site["photo"],
        "address": {"@type": "PostalAddress", "addressLocality": "Yogyakarta",
                    "addressCountry": "ID"},
        "alumniOf": [{"@type": "CollegeOrUniversity", "name": e["institution"]}
                     for e in d["education"]],
        "knowsLanguage": [pick(lg["name"], "en") for lg in d["languages"]],
        "sameAs": list(meta["links"].values()),
    }
    if meta.get("phone"):
        payload["telephone"] = meta["phone"]
    # Controlled data, but never let a '<' close the script element.
    return json.dumps(payload, ensure_ascii=False).replace("<", "\\u003c")


def _env():
    from jinja2 import Environment, FileSystemLoader

    # autoescape=True, not select_autoescape(["html"]): that helper matches on
    # the final extension, and these templates end in ".j2", so escaping was
    # silently off. A stray "&" or "<" in cv-data.yaml would reach the page raw.
    env = Environment(loader=FileSystemLoader(ROOT / "templates"),
                      autoescape=True)
    env.filters["runs_html"] = _runs_html
    # The portfolio ladder spans 35 -> 83,529, so bars are drawn on a log
    # scale. Deriving the width from the value keeps the chart honest: change
    # the number in the YAML and the bar moves with it.
    env.filters["log10"] = math.log10
    return env


def render_landing(d: dict, sections: dict, path: Path) -> None:
    from markupsafe import Markup

    clients = landing_clients(d)
    talks = [t for t in sections["talks"] if t["tier"] == 1][:5]
    path.write_text(_env().get_template("index.html.j2").render(
        meta=d["meta"],
        site=d["site"],
        clients=clients,
        work=landing_work(d),
        publications=sections["publications"],
        talks=talks,
        stats=landing_stats(d, clients),
        jsonld=Markup(landing_jsonld(d)),
        langs=LANGS,
        default_lang="id",
        year=dt.date.today().year,
    ), encoding="utf-8")


# --------------------------------------------------------------------------- #
# Web CV  (kristonova.github.io/cv/ — EN/ID toggle, Résumé/Full toggle, prints)
# --------------------------------------------------------------------------- #
def render_html(meta: dict, sections: dict, site: dict, path: Path) -> None:
    payload = [{"key": k, "title": L[k],
                "entries": visible(sections.get(k, []), PROFILES["academic"])}
               for k in ORDER["academic"] if sections.get(k)]
    # Sections the resume profile drops — so the page's "Résumé" toggle shows
    # exactly what resume.*.docx contains, not merely tier-1 of everything.
    academic_only = [k for k in ORDER["academic"] if k not in ORDER["resume"]]
    path.write_text(_env().get_template("cv.html.j2").render(
        meta=meta, sections=payload, labels=L, langs=LANGS, site=site,
        academic_only=academic_only), encoding="utf-8")


def _runs_html(runs: list) -> str:
    from markupsafe import Markup, escape
    out = []
    for r in runs:
        t = escape(r["t"])
        if r["b"]:
            t = Markup(f"<strong>{t}</strong>")
        if r["i"]:
            t = Markup(f"<em>{t}</em>")
        out.append(t)
    return Markup("").join(out)


# --------------------------------------------------------------------------- #
def report_todos() -> int:
    todos = [(i, ln.strip()) for i, ln in
             enumerate(DATA.read_text(encoding="utf-8").splitlines(), 1)
             if "TODO(user)" in ln]
    if todos:
        print(f"\n  {len(todos)} TODO(user) left in cv-data.yaml:")
        for i, ln in todos:
            print(f"    cv-data.yaml:{i}  {ln.lstrip('# ')}")
    return len(todos)


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--profile", choices=list(PROFILES), action="append")
    ap.add_argument("--lang", choices=list(LANGS), action="append")
    ap.add_argument("--format", choices=["md", "docx", "html"], action="append")
    a = ap.parse_args()
    profiles = a.profile or list(PROFILES)
    langs = a.lang or list(LANGS)
    formats = a.format or ["md", "docx", "html"]

    data = yaml.safe_load(DATA.read_text(encoding="utf-8"))
    meta, sections = data["meta"], build_model(data)
    OUT.mkdir(exist_ok=True)

    made = []
    for profile in profiles:
        stem = "resume" if profile == "resume" else "cv-academic"
        for lang in langs:
            if "md" in formats:
                p = OUT / f"{stem}.{lang}.md"
                p.write_text(render_markdown(meta, sections, lang, profile), encoding="utf-8")
                made.append(p)
            if "docx" in formats:
                p = OUT / f"{stem}.{lang}.docx"
                render_docx(meta, sections, lang, profile, p)
                made.append(p)
    if "html" in formats:
        render_html(meta, sections, data["site"], WEB_CV)
        made.append(WEB_CV)
        render_landing(data, sections, LANDING)
        made.append(LANDING)
        # Superseded by cv/index.html, which GitHub Pages serves at /cv/.
        stale = OUT / "cv.html"
        if stale.exists():
            stale.unlink()

    for p in made:
        print(f"  wrote {p.relative_to(ROOT.parent)}  ({p.stat().st_size:,} bytes)")
    print(f"\n  {len(made)} files generated.")
    report_todos()
    return 0


if __name__ == "__main__":
    sys.exit(main())
